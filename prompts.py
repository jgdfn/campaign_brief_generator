"""Exports the brief to a Word document, using python-docx."""

import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from schema import BRIEF_FIELDS


def _add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    doc.add_paragraph()


def build_docx_bytes(response_json: dict, campaign_title: str = "Campaign Brief") -> bytes:
    brief = response_json["brief"]
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    _add_heading(doc, campaign_title)

    for field in BRIEF_FIELDS:
        entry = brief.get(field, {"value": "Missing"})
        p_label = doc.add_paragraph()
        run_label = p_label.add_run(field)
        run_label.bold = True
        run_label.font.size = Pt(11)

        p_value = doc.add_paragraph()
        run_value = p_value.add_run(str(entry.get("value", "")))
        run_value.font.size = Pt(10.5)
        if entry.get("source_type") == "Missing":
            run_value.italic = True
            run_value.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
